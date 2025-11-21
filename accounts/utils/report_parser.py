# accounts/utils/report_parser.py
import re
from datetime import datetime
from io import BytesIO

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

MONTHS = {
    'january':1,'jan':1,'february':2,'feb':2,'march':3,'mar':3,'april':4,'apr':4,
    'may':5,'june':6,'jun':6,'july':7,'jul':7,'august':8,'aug':8,'september':9,'sep':9,'sept':9,
    'october':10,'oct':10,'november':11,'nov':11,'december':12,'dec':12
}

# Short code → full department name
DEPARTMENT_NICKNAMES = {
    "DCS (IT-CS)": "DEPARTMENT OF COMPUTER STUDIES",
    "DMS (HM)": "DEPARTMENT OF MANAGEMENT STUDIES (HM)",
    "DMS (BM)": "DEPARTMENT OF MANAGEMENT STUDIES (BM)",
    "DC (Criminology)": "DEPARTMENT OF CRIMINOLOGY",
    "DAS (Psychology)": "DEPARTMENT OF ARTS AND SCIENCES",
    "DTE (Education)": "DEPARTMENT OF TEACHER EDUCATION",
}
CHECK_MARKS = ['yes', 'y', '✓', '✔', '☑', '■', '●', 'x', '✗', 'checked']

def is_checked(value):
    if not value:
        return False
    v = value.lower().strip()

    # exact YES should count
    if v == "yes":
        return True

    # exact symbols should count
    if any(sym in value for sym in ['✔', '✓', '☑']):
        return True

    # ignore decorative boxes
    if any(sym in value for sym in ['■', '□', '☐', '◻', '◼']):
        return False

    return False


def safe_read_bytes(fieldfile):
    try:
        fieldfile.seek(0)
    except Exception:
        pass
    return fieldfile.read()

# ---------- TEXT + TABLE EXTRACTION ----------
def extract_text_from_file(fieldfile):
    bf = safe_read_bytes(fieldfile)
    lines, table_kv = [], {}

    # DOCX
    if docx:
        try:
            document = docx.Document(BytesIO(bf))
            for p in document.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2 and cells[0] and cells[1]:
                        table_kv[cells[0]] = cells[1]
                    for c in cells:
                        if c:
                            for sub in c.splitlines():
                                s = sub.strip()
                                if s:
                                    lines.append(s)
            return lines, table_kv
        except Exception:
            pass

    # PDF (improved handling)
    if bf[:4] == b"%PDF" and pdfplumber:
        try:
            with pdfplumber.open(BytesIO(bf)) as pdf:
                for page in pdf.pages:
                    # Extract text with layout-aware mode
                    txt = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                    txt = re.sub(r'[ \t]+', ' ', txt)  # normalize spaces
                    for l in txt.splitlines():
                        l = l.strip()
                        if l:
                            lines.append(l)
            # Attempt to reconstruct "Key: Value" pairs
            for i, line in enumerate(lines):
                # Combine cases like "Title of Training" on one line and value on next
                if re.search(r':$', line) and i + 1 < len(lines):
                    key = line.rstrip(':').strip()
                    val = lines[i + 1].strip()
                    if key and val:
                        table_kv[key] = val
                elif ':' in line:
                    key, _, val = line.partition(':')
                    if key.strip() and val.strip():
                        table_kv[key.strip()] = val.strip()
            return lines, table_kv
        except Exception as e:
            print("PDF parse error:", e)

    # fallback text
    try:
        txt = bf.decode("utf-8", errors="ignore")
        return [l.strip() for l in txt.splitlines() if l.strip()], {}
    except Exception:
        return [], {}


# ---------- HELPERS ----------
def normalize_key(k):
    return re.sub(r'\W+', '', k.lower()) if k else ''

def to_int(val):
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    s = s.replace(',', '').replace(' ', '')
    if re.match(r'^(n/?a|-|na|none|nil)$', s, re.I):
        return None
    try:
        return int(float(s))
    except Exception:
        return None

KEY_ALIASES = {
    'title': ['titleoftraining','trainingtitle','title','titleofactivity','titleoftrainingactivity','titleof the training','titleoftraining:'],
    'date': ['dateconducted','date','inclusive dates','inclusive','dates','dateconducted:','date conducted','date of training'],
    'days': ['noofdays','numberofdays','nodays','no.days','noofdays:','no. of days','duration'],
    'male': ['male','maleparticipants'],
    'female': ['female','femaleparticipants'],
    'student': ['student','students'],
    'farmer': ['farmer','farmers'],
    'fisherfolk': ['fisherfolk','fisherman','fishermen'],
    'agricultural_technician': ['agriculturaltechnician','agriculturaltech','agritech','agriculturaltechnician'],
    'government_employee': ['governmentemployee','govemployee','governmentemployee'],
    'private_employee': ['privateemployee','privemployee','privateemployees'],
    'other_category': ['others','other','othercategory'],
    'total_by_category': ['total','totalparticipants','totalbycategory','totalparticipantsbycategory'],
    'surveyed': ['noofparticipantssurveyed','totaltraineessurveyed','surveyed','totaltraineessurveyed'],
    'tvl_solo_parent': [
        'soloparent',
        'solo parent',
        'noofparticipantssoloparent',
        'noofparticipantswhorear.soloparents',
        'noofparticipantswhorear.soloparent',
        'no.ofparticipantswhorear.soloparents',
        'no.ofparticipantswhorear.soloparent',
        'noofparticipantswhoaresoloparents',
        'noofparticipantswhoaresoloparent',
        'no.ofparticipantswhoaresoloparents',
        'no.ofparticipantswhoaresoloparent',
        'noofparticipantswhoaresoloparent(s)',
        'no.ofparticipantswhoaresoloparent(s)',
        'noofparticipantswhosoloparent',
        'no.ofparticipantswhosoloparent',
        'no. of participants who are solo parents',
        'no. of participants who are solo parent',
    ],

    'tvl_pwd': [
        'pwd',
        'personswithdisability',
        'participantswithdisability',
        'noofparticipantspwd',
        'no.ofparticipantspwd',
        'noofparticipantswhopwd',
        'no.ofparticipantswhopwd',
        'noofparticipantswithdisabilities',
        'no.ofparticipantswithdisabilities',
        'noofparticipantswithdisability',
        'no.ofparticipantswithdisability',
        'no. of participants with disabilities',
        'no. of participants with disability',
    ],
    'tvl_4ps': ['4ps','4-ps','fourps','noofparticipants4ps'],
    'tvl_pwd_type': ['typeofdisability','pwddescription'],
    'collaborating_agencies': ['collaboratingagencies','collaboratingagency','partneragencies','partneragency','partneragencyies'],
    'amount_charged_to_cvsu': [
        'amountchargedtocvsu',
        'amountchargedtocvsucampuscollegeunit',
        'amountchargedtocvsu(campuscollegeunit)',
        'amountchargedtocvsu(campus/college/unit)',
        'amountchargedtocampus',
        'amountchargedtounit',
    ],
    'amount_charged_to_partner_agency': [
        'amountchargedtopartneragency',
        'amountchargedtopartneragencyphp',
        'amountchargedtopartneragencyphpifthereisnocashinvolveincludeestimatesvalue',
        'amountchargedtopartneragencyifthereisnocashinvolveincludeestimatesvalue',
        'amountchargedtopartneragencyphpifthereisnocashinvolveincludeestimatevalue',
    ],

    'category_of_training': [
    'categoryoftraining',
    'categorytraining',
    'trainingcategory',
    'category'
],


    'venue': ['place','venue','location','trainingvenue','traininglocation'],
}


def get_value(kv, field):
    for alias in KEY_ALIASES.get(field, []):
        if alias in kv and kv[alias] not in (None, ''):
            return kv[alias]
    return None

# ---------- DATE HELPERS ----------
def parse_date_range(text):
    if not text:
        return None, None, ""
    text = re.sub(r'\s*[-–—]\s*', '-', text)
    text = re.sub(r'\s*(to)\s*', '-', text, flags=re.I)

    pat1 = re.compile(r'(?P<mon>\b(?:' + '|'.join(MONTHS.keys()) + r')\b)\s*(?P<d1>[0-3]?\d)(?:-(?P<d2>[0-3]?\d))?,?\s*(?P<yr>\d{4})', re.I)
    m = pat1.search(text)
    if m:
        mon = MONTHS.get(m.group('mon').lower(),1)
        d1 = int(m.group('d1'))
        d2 = int(m.group('d2')) if m.group('d2') else d1
        yr = int(m.group('yr'))
        try: start,end = datetime(yr,mon,d1).date(), datetime(yr,mon,d2).date()
        except: start,end = None,None
        return start,end,m.group(0)

    pat2 = re.compile(r'(?P<d1>[0-3]?\d)(?:-(?P<d2>[0-3]?\d))?\s*(?P<mon>\b(?:' + '|'.join(MONTHS.keys()) + r')\b),?\s*(?P<yr>\d{4})', re.I)
    m2 = pat2.search(text)
    if m2:
        mon = MONTHS.get(m2.group('mon').lower(),1)
        d1 = int(m2.group('d1'))
        d2 = int(m2.group('d2')) if m2.group('d2') else d1
        yr = int(m2.group('yr'))
        try: start,end = datetime(yr,mon,d1).date(), datetime(yr,mon,d2).date()
        except: start,end = None,None
        return start,end,m2.group(0)

    pat3 = re.compile(r'(?P<mon>\b(?:' + '|'.join(MONTHS.keys()) + r')\b)\s*(?P<d>[0-3]?\d),?\s*(?P<yr>\d{4})', re.I)
    m3 = pat3.search(text)
    if m3:
        mon = MONTHS.get(m3.group('mon').lower(),1)
        d = int(m3.group('d'))
        yr = int(m3.group('yr'))
        try: dt = datetime(yr,mon,d).date()
        except: dt = None
        return dt, dt, m3.group(0)

    return None,None,""

def days_between(start,end):
    if not start or not end:
        return None
    try:
        return max(1,(end-start).days+1)
    except: return None

def weighting_multiplier_from_days(days, total_by_category):
    """
    Compute the weighting multiplier based strictly on numeric training days.
    If days is not 5, 4, 3, 2, or 1 → multiplier = 0.5.
    """

    # Safely convert total_by_category
    try:
        total_by_category = int(total_by_category) if total_by_category is not None else 0
    except (ValueError, TypeError):
        total_by_category = 0

    # Convert days to a number if possible
    try:
        days = float(days)
    except Exception:
        days = 0

    # Multiplier rules
    multiplier_map = {
        5: 2.0,
        4: 1.5,
        3: 1.5,
        2: 1.25,
        1: 1.0,
    }

    # If days matches exact known rule → use it
    # Otherwise → default 0.5
    multiplier = multiplier_map.get(days, 0.5)

    return total_by_category * multiplier



# ---------- RATINGS ----------
def parse_rating_block(lines,start_index):
    counts = {}
    i = start_index+1
    tokens=[]
    limit=30
    while i<len(lines) and len(tokens)<limit:
        token=lines[i].strip()
        if not token: 
            i+=1
            continue
        if re.search(r'[A-Za-z]', token) and not re.search(r'\d', token):
            break
        tokens.append(token)
        i += 1

    flat = []
    for t in tokens:
        if ':' in t:
            left,right = [x.strip() for x in t.split(':',1)]
            flat.extend([left,right])
        elif '-' in t and re.search(r'\d',t):
            left,right = [x.strip() for x in t.split('-',1)]
            flat.extend([left,right])
        else:
            flat.extend(t.split())

    j=0
    while j<len(flat)-1:
        a,b=flat[j],flat[j+1]
        if re.match(r'^[1-5]$',a) and re.match(r'^[0-9,]+$',b):
            counts[a]=int(b.replace(',',''))
            j+=2
            continue
        if re.match(r'^[0-9,]+$',a) and re.match(r'^[1-5]$',b):
            counts[b]=int(a.replace(',',''))
            j+=2
            continue
        j+=1
    return counts,i

def compute_average_from_counts(counts_dict,total_surveyed):
    if not counts_dict or not total_surveyed:
        return None
    try:
        s=sum(int(k)*int(v) for k,v in counts_dict.items())
        return round(s/total_surveyed,2) if total_surveyed else None
    except:
        return None

# ---------- MAIN PARSER ----------
def parse_report(fieldfile, uploader_user=None):
    lines, table_kv = extract_text_from_file(fieldfile)
    kv = {normalize_key(k): v for k, v in table_kv.items() if v}

    # Parse "Key: Value" lines from text
    for idx, l in enumerate(lines):
        if ':' in l:
            key, _, val = l.partition(':')
            if key:
                key_norm = normalize_key(key)
                val = val.strip()
                if not val and idx + 1 < len(lines):
                    val = lines[idx + 1].strip()
                kv[key_norm] = val

    # ---------- uploader info ----------
    department = getattr(uploader_user, 'department', '') if uploader_user else ''
    department = DEPARTMENT_NICKNAMES.get(department, department)

    contact_person = getattr(uploader_user, 'full_name', None) or getattr(uploader_user, 'get_full_name', None)
    if callable(contact_person):
        try:
            contact_person = contact_person()
        except:
            contact_person = ''
    contact_person = contact_person or getattr(uploader_user, 'username', '') or getattr(uploader_user, 'email', '')
    number_email = getattr(uploader_user, 'email', '') if uploader_user else ''

    # ---------- TITLE ----------
    title = get_value(kv, 'title') or ''
    if not title:
        for k, v in kv.items():
            if any(alias in k for alias in KEY_ALIASES['title']):
                title = v.strip()
                break
    if not title:
        for idx, l in enumerate(lines[:50]):
            if ':' in l:
                key, _, val = l.partition(':')
                if normalize_key(key) in KEY_ALIASES['title']:
                    title = val.strip() if val.strip() else (lines[idx + 1].strip() if idx + 1 < len(lines) else '')
                    break
            elif l.lower().startswith('title'):
                _, _, val = l.partition(':')
                title = val.strip() if val.strip() else (lines[idx + 1].strip() if idx + 1 < len(lines) else '')
                if title:
                    break

    # ---------- DATES & DAYS ----------
    date_text = get_value(kv, 'date') or ''
    days = to_int(get_value(kv, 'days')) or 0
    if not date_text or not days:
        for idx, l in enumerate(lines[:50]):
            key, _, val = l.partition(':')
            norm_key = normalize_key(key)
            if not date_text and norm_key in KEY_ALIASES['date']:
                date_text = (val or (lines[idx + 1] if idx + 1 < len(lines) else '')).strip()
            if not days and any(alias in norm_key for alias in KEY_ALIASES['days']):
                m = re.search(r'\d+', val or (lines[idx + 1] if idx + 1 < len(lines) else ''))
                if m:
                    days = int(m.group(0))

    start_date, end_date, raw_match = parse_date_range(date_text)

    # ---------- PARTICIPANTS ----------
    male = to_int(get_value(kv, 'male')) or 0
    female = to_int(get_value(kv, 'female')) or 0
    student = to_int(get_value(kv, 'student')) or 0
    farmer = to_int(get_value(kv, 'farmer')) or 0
    fisherfolk = to_int(get_value(kv, 'fisherfolk')) or 0
    agri_tech = to_int(get_value(kv, 'agricultural_technician')) or 0
    gov_emp = to_int(get_value(kv, 'government_employee')) or 0
    priv_emp = to_int(get_value(kv, 'private_employee')) or 0
    other_cat = to_int(get_value(kv, 'other_category')) or 0
    total_by_category = to_int(get_value(kv, 'total_by_category')) or 0

    # Calculate total participants reliably
    numbers = [
        total_by_category,
        male + female,
        student + farmer + fisherfolk + agri_tech + gov_emp + priv_emp + other_cat
    ]
    total_participants = next((n for n in numbers if n > 0), 0)

    # ---------- NEW FIELD: total_persons_trained ----------
    total_persons_trained = total_by_category or total_participants or 0

    total_surveyed = to_int(get_value(kv, 'surveyed')) or 0

    # ---------- RATINGS ----------
    relevance_counts, quality_counts, timeliness_counts = {}, {}, {}
    for idx, ln in enumerate(lines):
        low = ln.lower()
        if 'relevance' in low and not relevance_counts:
            relevance_counts, _ = parse_rating_block(lines, idx)
        if 'quality' in low and not quality_counts:
            quality_counts, _ = parse_rating_block(lines, idx)
        if 'timeliness' in low and not timeliness_counts:
            timeliness_counts, _ = parse_rating_block(lines, idx)

    relevance_average = compute_average_from_counts(relevance_counts, total_surveyed)
    quality_average = compute_average_from_counts(quality_counts, total_surveyed)
    timeliness_average = compute_average_from_counts(timeliness_counts, total_surveyed)
    # ---------- CATEGORY OF TRAINING FIX ----------
    category = ''
    CATEGORY_OPTIONS = {
        "TVL": "technical",
        "AE": "agricultural",
        "CE": "continuing",
        "BE": "basic education",
        "GAD": "gender",
        "O": "others",
    }

    TRUE_MARKS = {"yes", "✔", "✓", "☑"}

    def is_checked_line(line):
        return line and any(sym in line.lower() for sym in TRUE_MARKS)

    def extract_category_from_line(line):
        """Return matching category code if line contains a keyword."""
        if not line:
            return None
        low = line.lower()
        for code, keyword in CATEGORY_OPTIONS.items():
            if keyword in low:
                return code
        return None

    # locate "Category of Training" block
    start_index = None
    for i, line in enumerate(lines):
        if "category of training" in line.lower():
            start_index = i
            break

    if start_index is not None:
        block = lines[start_index:start_index + 20]  # read next 20 lines

        for idx, line in enumerate(block):
            if is_checked_line(line):
                # 1️⃣ Check same line
                code = extract_category_from_line(line)
                if code:
                    category = code
                    break

                # 2️⃣ Check next non-empty line
                next_line = ''
                for j in range(idx + 1, len(block)):
                    if block[j].strip():
                        next_line = block[j].strip()
                        break
                code = extract_category_from_line(next_line)
                if code:
                    category = code
                    break

    # ---------- TVL ----------
    def tvl_int(key):
        val = get_value(kv, key)
        if val is None:
            return None
        s = str(val).strip()
        if re.match(r'^(n/?a|-|na|none)$', s, re.I):
            return None
        return to_int(s)

    tvl_solo_parent = tvl_int('tvl_solo_parent')
    tvl_4ps = tvl_int('tvl_4ps')
    tvl_pwd = tvl_int('tvl_pwd')
    tvl_pwd_type = get_value(kv, 'tvl_pwd_type') or ''
    if tvl_pwd_type and re.match(r'^(n/?a|-|na|none)$', tvl_pwd_type, re.I):
        tvl_pwd_type = ''

    # ---------- WEIGHT ----------
    weight_multiplier = weighting_multiplier_from_days(days, total_by_category)
    weighted_persons = round(total_participants * weight_multiplier, 2) if weight_multiplier else 0

    # ---------- NEW FIELDS ----------
    collaborating_agencies = get_value(kv, 'collaborating_agencies') or ''
    amount_charged_to_cvsu = get_value(kv, 'amount_charged_to_cvsu') or '0'
    amount_charged_to_partner_agency = get_value(kv, 'amount_charged_to_partner_agency') or '0'
    venue = get_value(kv, 'venue') or ''


    return {
        'department': department,
        'contact_person': contact_person,
        'number_email': number_email,
        'project_no': '',
        'category': category,
        'title': title,
        'date_conducted_text': date_text,
        'start_date': start_date,
        'end_date': end_date,
        'number_of_days': days,
        'male_participants': male,
        'female_participants': female,
        'total_participants': total_participants,
        'student': student,
        'farmer': farmer,
        'fisherfolk': fisherfolk,
        'agricultural_technician': agri_tech,
        'government_employee': gov_emp,
        'private_employee': priv_emp,
        'other_category': other_cat,
        'total_by_category': total_by_category,
        'total_persons_trained': total_persons_trained,
        'tvl_solo_parent': tvl_solo_parent,
        'tvl_4ps': tvl_4ps,
        'tvl_pwd': tvl_pwd,
        'tvl_pwd_type': tvl_pwd_type,
        'total_trainees_surveyed': total_surveyed,
        'relevance_counts': relevance_counts,
        'relevance_average': relevance_average,
        'quality_counts': quality_counts,
        'quality_average': quality_average,
        'timeliness_counts': timeliness_counts,
        'timeliness_average': timeliness_average,
        'weight_multiplier': weight_multiplier,
        'weighted_persons': weighted_persons,
        'collaborating_agencies': collaborating_agencies,
        'amount_charged_to_cvsu': amount_charged_to_cvsu,
        'amount_charged_to_partner_agency': amount_charged_to_partner_agency,
        'venue': venue,
    }
